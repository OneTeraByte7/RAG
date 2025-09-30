"""
Document processing for PDF and DOCX files
"""
import pdfplumber
from docx import Document
from pathlib import Path
from typing import List, Dict, Optional
from loguru import logger
import hashlib
from datetime import datetime

from config.settings import settings


class DocumentChunker:
    """Chunk documents into semantic segments"""
    
    def __init__(
        self,
        chunk_size: int = None,
        chunk_overlap: int = None
    ):
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
    
    def chunk_text(self, text: str, metadata: Dict = None) -> List[Dict]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Text to chunk
            metadata: Additional metadata to attach to each chunk
            
        Returns:
            List of chunk dicts
        """
        # Simple word-based chunking
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            
            chunk_data = {
                "text": chunk_text,
                "chunk_id": i // (self.chunk_size - self.chunk_overlap),
                "start_word": i,
                "end_word": min(i + self.chunk_size, len(words)),
                "metadata": metadata or {}
            }
            
            chunks.append(chunk_data)
        
        return chunks


class PDFProcessor:
    """Process PDF documents"""
    
    def __init__(self):
        self.chunker = DocumentChunker()
        
    def extract_text(self, pdf_path: str) -> Dict:
        """
        Extract text and metadata from PDF
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Dict with text, pages, and metadata
        """
        logger.info(f"Processing PDF: {pdf_path}")
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                pages_data = []
                full_text = []
                
                for page_num, page in enumerate(pdf.pages, start=1):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        full_text.append(text)
                        
                        # Extract tables
                        tables = page.extract_tables()
                        
                        # Extract images
                        images = page.images
                        
                        pages_data.append({
                            "page_number": page_num,
                            "text": text,
                            "num_tables": len(tables) if tables else 0,
                            "num_images": len(images) if images else 0,
                            "tables": tables if tables else []
                        })
                
                # Generate document ID
                doc_id = self._generate_doc_id(pdf_path)
                
                result = {
                    "doc_id": doc_id,
                    "source": Path(pdf_path).name,
                    "full_text": "\n\n".join(full_text),
                    "pages": pages_data,
                    "total_pages": len(pdf.pages),
                    "metadata": {
                        "file_path": str(pdf_path),
                        "file_size": Path(pdf_path).stat().st_size,
                        "processed_at": datetime.now().isoformat()
                    }
                }
                
                logger.info(f"PDF processed: {len(pdf.pages)} pages")
                return result
                
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path}: {e}")
            raise
    
    def process_and_chunk(self, pdf_path: str) -> List[Dict]:
        """
        Extract and chunk PDF content
        
        Returns:
            List of chunks ready for embedding
        """
        # Extract content
        doc_data = self.extract_text(pdf_path)
        
        # Chunk by pages
        all_chunks = []
        for page in doc_data["pages"]:
            page_metadata = {
                "doc_id": doc_data["doc_id"],
                "source": doc_data["source"],
                "page": page["page_number"],
                "type": "pdf"
            }
            
            # Chunk page text
            chunks = self.chunker.chunk_text(page["text"], page_metadata)
            all_chunks.extend(chunks)
        
        logger.info(f"Created {len(all_chunks)} chunks from PDF")
        return all_chunks
    
    def _generate_doc_id(self, file_path: str) -> str:
        """Generate unique document ID"""
        return hashlib.md5(str(file_path).encode()).hexdigest()[:16]


class DOCXProcessor:
    """Process DOCX documents"""
    
    def __init__(self):
        self.chunker = DocumentChunker()
    
    def extract_text(self, docx_path: str) -> Dict:
        """
        Extract text and metadata from DOCX
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Dict with text, paragraphs, and metadata
        """
        logger.info(f"Processing DOCX: {docx_path}")
        
        try:
            doc = Document(docx_path)
            
            # Extract paragraphs with style info
            paragraphs = []
            full_text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    para_data = {
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal",
                        "is_heading": para.style.name.startswith("Heading") if para.style else False
                    }
                    paragraphs.append(para_data)
                    full_text.append(para.text)
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Generate document ID
            doc_id = self._generate_doc_id(docx_path)
            
            result = {
                "doc_id": doc_id,
                "source": Path(docx_path).name,
                "full_text": "\n\n".join(full_text),
                "paragraphs": paragraphs,
                "tables": tables,
                "metadata": {
                    "file_path": str(docx_path),
                    "file_size": Path(docx_path).stat().st_size,
                    "num_paragraphs": len(paragraphs),
                    "num_tables": len(tables),
                    "processed_at": datetime.now().isoformat()
                }
            }
            
            logger.info(f"DOCX processed: {len(paragraphs)} paragraphs")
            return result
            
        except Exception as e:
            logger.error(f"Error processing DOCX {docx_path}: {e}")
            raise
    
    def process_and_chunk(self, docx_path: str) -> List[Dict]:
        """
        Extract and chunk DOCX content
        
        Returns:
            List of chunks ready for embedding
        """
        # Extract content
        doc_data = self.extract_text(docx_path)
        
        # Chunk paragraphs
        metadata = {
            "doc_id": doc_data["doc_id"],
            "source": doc_data["source"],
            "type": "docx"
        }
        
        chunks = self.chunker.chunk_text(doc_data["full_text"], metadata)
        
        logger.info(f"Created {len(chunks)} chunks from DOCX")
        return chunks
    
    def _generate_doc_id(self, file_path: str) -> str:
        """Generate unique document ID"""
        return hashlib.md5(str(file_path).encode()).hexdigest()[:16]


class DocumentProcessor:
    """Unified document processor"""
    
    def __init__(self):
        self.pdf_processor = PDFProcessor()
        self.docx_processor = DOCXProcessor()
    
    def process_document(self, file_path: str) -> List[Dict]:
        """
        Process document based on file type
        
        Args:
            file_path: Path to document
            
        Returns:
            List of chunks
        """
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.pdf':
            return self.pdf_processor.process_and_chunk(str(file_path))
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            return self.docx_processor.process_and_chunk(str(file_path))
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")